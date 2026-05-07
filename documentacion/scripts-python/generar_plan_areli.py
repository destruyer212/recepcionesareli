from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "Planificacion_Sistema_Recepciones_Areli.docx"


BRAND = RGBColor(115, 34, 78)
ACCENT = RGBColor(197, 151, 73)
DARK = RGBColor(42, 42, 42)
MUTED = RGBColor(95, 95, 95)
LIGHT = "F7F2F5"
GOLD_LIGHT = "FBF6EA"
HEADER_BG = "73224E"
HEADER_GOLD = "C59749"
TABLE_HEADER = "EFE5EB"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9D2D8", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
    p.style = style
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, title, text, fill=GOLD_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="E3D2A7")
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=BRAND)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, text, color=DARK)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False if widths else True
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BRAND
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Heading 1", 18, BRAND),
        ("Heading 2", 14, BRAND),
        ("Heading 3", 11.5, ACCENT),
    ]:
        st = styles[name]
        st.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), st.font.name)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(6)

    body = styles["Body Text"]
    body.font.name = "Aptos"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    body.font.size = Pt(10.5)
    body.font.color.rgb = DARK
    body.paragraph_format.line_spacing = 1.08
    body.paragraph_format.space_after = Pt(6)

    for list_style in ("List Bullet", "List Number"):
        st = styles[list_style]
        st.font.name = "Aptos"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        st.font.size = Pt(10.2)
        st.paragraph_format.space_after = Pt(2)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    add_run(p, "RECEPCIONES ARELI", bold=True, color=BRAND, size=26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Planificacion del Sistema Administrativo", bold=True, color=DARK, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Reservas, contratos, pagos, personal, gastos, inventario y balance financiero", color=MUTED, size=11)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell, color="E5D9E1")
    set_cell_margins(cell, top=260, start=260, bottom=260, end=260)
    p = cell.paragraphs[0]
    add_run(p, "Objetivo del documento", bold=True, color=BRAND, size=12)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.1
    p2.add_run(
        "Presentar una estructura clara para desarrollar un sistema que controle los eventos de Recepciones Areli en sus 3 pisos, "
        "desde la separacion de fecha hasta el cierre financiero del evento."
    )

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    rows = [
        ("Cliente", "Recepciones Areli"),
        ("Version", "Plan funcional inicial"),
        ("Fecha", "Abril 2026"),
        ("Alcance", "Sistema interno de administracion y control"),
    ]
    for i, (k, v) in enumerate(rows):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        for c in meta.row_cells(i):
            set_cell_border(c, color="E1DAE0")
            set_cell_margins(c, top=100, start=120, bottom=100, end=120)
        set_cell_shading(meta.cell(i, 0), TABLE_HEADER)
        meta.cell(i, 0).paragraphs[0].runs[0].bold = True
        meta.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = BRAND

    doc.add_page_break()


def main():
    doc = Document()
    style_document(doc)
    add_cover(doc)

    add_heading(doc, "1. Resumen ejecutivo", 1)
    add_body(
        doc,
        "El proyecto consiste en crear una plataforma para administrar Recepciones Areli de manera ordenada. "
        "El sistema controlara reservas, contratos, paquetes, pagos, garantias, organizacion del evento, personal, gastos, inventario y balance financiero."
    )
    add_body(
        doc,
        "Como el negocio trabajara con 3 pisos, el calendario debe permitir separar fechas y horarios por ambiente, evitando cruces dentro del mismo piso."
    )
    add_note(
        doc,
        "Idea central",
        "La primera version debe enfocarse en lo mas importante para el negocio: calendario por piso, contratos, pagos, garantias y balance basico."
    )

    add_heading(doc, "2. Objetivos del sistema", 1)
    add_bullets(
        doc,
        [
            "Controlar las reservas de eventos por fecha, horario y piso.",
            "Generar contratos con los datos reales del cliente, evento, piso y paquete contratado.",
            "Registrar pagos, adelantos, saldos y boletas internas.",
            "Manejar garantias de forma separada, incluyendo devoluciones o retenciones por daños.",
            "Organizar el personal asignado a cada evento y sus pagos.",
            "Registrar gastos del local y gastos asociados a eventos.",
            "Controlar el inventario del local, su ubicacion, estado y costo.",
            "Calcular balances por evento, por mes y por año.",
            "Guardar material visual de eventos para galeria, reels y futuras ventas.",
        ],
    )

    add_heading(doc, "3. Alcance general", 1)
    add_table(
        doc,
        ["Area", "Que controlara el sistema"],
        [
            ("Ventas y reservas", "Clientes, consultas, fechas separadas, pisos disponibles y estado de cada evento."),
            ("Contratos", "Contrato por piso, monto total, modalidad de pago, clausulas, capacidad y firmas."),
            ("Finanzas", "Pagos de clientes, boletas internas, garantias, gastos, pagos de personal y balance."),
            ("Operacion", "Checklist del evento, servicios incluidos, personal asignado y estado de preparacion."),
            ("Inventario", "Bienes del local, cantidad, costo, estado y ubicacion por piso o almacen."),
            ("Marketing", "Fotos, videos, reels, paquetes y material comercial."),
        ],
        widths=[4.2, 11.5],
    )

    add_heading(doc, "4. Pisos y ambientes", 1)
    add_body(
        doc,
        "El sistema trabajara con 3 pisos porque el tercer piso se encuentra en construccion. Cada piso debe manejarse como un ambiente independiente."
    )
    add_table(
        doc,
        ["Piso", "Uso en el sistema", "Datos sugeridos"],
        [
            ("1er piso", "Disponible para reservas y contratos.", "Capacidad, area, banos, cocina/lavadero, estado y observaciones."),
            ("2do piso", "Disponible para reservas y contratos.", "Capacidad, area, barra, cocina, banos, estado y observaciones."),
            ("3er piso", "Ambiente en construccion y futura reserva.", "Estado de obra, fecha estimada de uso, capacidad proyectada y restricciones."),
        ],
        widths=[3.0, 5.5, 7.0],
    )

    add_heading(doc, "5. Modulos del sistema", 1)
    modules = [
        ("Dashboard principal", "Resumen de eventos, pagos pendientes, reservas, gastos, ganancia estimada y ocupacion por piso."),
        ("Clientes", "Registro de datos, DNI/RUC, telefono, WhatsApp, direccion, observaciones e historial de eventos."),
        ("Calendario de reservas", "Vista diaria, semanal y mensual por piso, fecha, horario, estado y encargada responsable."),
        ("Paquetes", "Catalogo de paquetes como Basico, Areli, Premium, Promociones escolares, Matrimonios, 15 anos y personalizados."),
        ("Contratos", "Creacion de contratos por cliente, evento, piso, horario, monto, clausulas, capacidad y estado."),
        ("Pagos de clientes", "Adelantos, pagos parciales, saldo pendiente, medio de pago, responsable y observaciones."),
        ("Garantias", "Recepcion, devolucion, retencion parcial o total por desperfectos y motivo del descuento."),
        ("Boletas internas", "Comprobante interno por cada pago recibido, con numero, concepto, monto y saldo restante."),
        ("Organizacion del evento", "Checklist, servicios incluidos, servicios adicionales, horarios y requerimientos especiales."),
        ("Personal", "Registro de encargada, coordinadora, mozos, barman, DJ, animador, seguridad, limpieza, chef y otros roles."),
        ("Pago de personal", "Monto acordado por evento, adelantos, saldo, fecha y estado de pago."),
        ("Gastos del local", "Agua, luz, SUNAT, impuestos, APDAYC, mantenimiento, publicidad, reparaciones, compras y otros."),
        ("Inventario", "Mesas, sillas, vajilla, copas, luces, parlantes, decoracion, muebles, equipos y su costo."),
        ("Galeria y reels", "Fotos, videos, reels, tipo de evento, piso, fecha y material visible o privado."),
        ("Balance financiero", "Ingresos, egresos, pagos pendientes, garantias, gastos, utilidad por evento y reportes mensuales."),
        ("Usuarios y permisos", "Acceso por rol: administrador, encargada, caja/finanzas, marketing y personal."),
    ]
    add_table(doc, ["Modulo", "Descripcion"], modules, widths=[4.4, 11.3])

    add_heading(doc, "6. Estados principales", 1)
    add_table(
        doc,
        ["Elemento", "Estados sugeridos"],
        [
            ("Evento", "Consulta, Separado, Contratado, En preparacion, Realizado, Cerrado, Cancelado, Reprogramado."),
            ("Contrato", "Borrador, Firmado, Anulado, Finalizado."),
            ("Pago del cliente", "Pendiente, Parcial, Cancelado."),
            ("Garantia", "Pendiente, Recibida, Devuelta, Retenida parcialmente, Retenida totalmente."),
            ("Pago de personal", "Pendiente, Pagado parcial, Pagado."),
            ("Inventario", "Bueno, Regular, Danado, Perdido, En reparacion."),
        ],
        widths=[4.0, 11.7],
    )

    add_heading(doc, "7. Flujo completo de un evento", 1)
    add_numbered(
        doc,
        [
            "El cliente consulta disponibilidad.",
            "Se revisa el calendario por piso.",
            "Se registra al cliente.",
            "Se separa fecha, piso y horario.",
            "Se elige el paquete o se arma un paquete personalizado.",
            "Se registra adelanto o separacion.",
            "Se genera la boleta interna del pago.",
            "Se crea el contrato con los datos del evento.",
            "Se firma o confirma el contrato.",
            "Se organiza personal y servicios incluidos.",
            "Se registran pagos restantes.",
            "Se recibe garantia, si corresponde.",
            "Se realiza el evento.",
            "Se revisan danos, desperfectos o pendientes.",
            "Se devuelve o retiene garantia segun corresponda.",
            "Se paga al personal asignado.",
            "Se registran gastos del evento.",
            "Se cierra el evento.",
            "El sistema calcula el balance final.",
        ],
    )

    add_heading(doc, "8. Reglas del negocio", 1)
    add_bullets(
        doc,
        [
            "No se debe permitir reservar el mismo piso en la misma fecha y horario para dos eventos distintos.",
            "Si los eventos son en pisos diferentes, el sistema puede permitirlos siempre que el negocio lo autorice.",
            "La garantia debe registrarse separada del ingreso del evento porque puede devolverse.",
            "Cada pago debe quedar asociado a un evento y generar una boleta interna.",
            "El saldo pendiente debe calcularse automaticamente.",
            "El contrato debe guardar el piso contratado: 1er, 2do o 3er piso.",
            "Los paquetes deben permitir servicios incluidos y servicios adicionales.",
            "El balance debe considerar ingresos, pagos al personal y gastos asociados.",
            "Un evento no deberia cerrarse si existen pagos pendientes, personal sin pagar o garantia sin resolver.",
        ],
    )

    add_heading(doc, "9. Paquetes y servicios", 1)
    add_body(
        doc,
        "A partir de los documentos revisados, el sistema debe permitir registrar paquetes con precios base y condiciones. Estos valores deben poder editarse porque el negocio puede cambiarlos con el tiempo."
    )
    add_table(
        doc,
        ["Paquete / categoria", "Datos a registrar"],
        [
            ("Paquete Basico", "Precio base, capacidad, local, mobiliario, decoracion basica, sonido, garantia y porcentaje de separacion."),
            ("Paquete Areli", "Servicios ampliados: mesas, menaje, torta, bocaditos, cena, brindis, sonido, foto/video, maestro de ceremonia y personal."),
            ("Paquete Premium", "Servicios completos: cena, brindis, hora loca, barman, candy bar, foto/video, personal y condiciones."),
            ("Promociones escolares", "Precio por alumno, cantidad de alumnos, dia de semana, cubierto adicional, separacion, bar y condiciones especiales."),
            ("Personalizados", "Paquetes armados a medida con servicios seleccionados y precio ajustado."),
        ],
        widths=[4.4, 11.3],
    )

    add_heading(doc, "10. Datos clave por contrato", 1)
    add_bullets(
        doc,
        [
            "Numero de contrato.",
            "Datos del arrendador o representante.",
            "Datos del cliente o arrendatario.",
            "DNI/RUC del cliente.",
            "Direccion del cliente.",
            "Tipo de evento.",
            "Fecha del evento.",
            "Hora de inicio y hora de finalizacion.",
            "Piso contratado: 1er, 2do o 3er piso.",
            "Capacidad maxima.",
            "Paquete contratado.",
            "Monto total.",
            "Modalidad de pago.",
            "Adelanto y saldo pendiente.",
            "Garantia.",
            "Clausulas de responsabilidad, seguridad, menores de edad, desperfectos y entrega del local.",
            "Firmas de arrendador y arrendatario.",
        ],
    )

    add_heading(doc, "11. Balance financiero", 1)
    add_body(doc, "El balance debe permitir medir si cada evento y cada mes fueron rentables.")
    add_table(
        doc,
        ["Reporte", "Que debe mostrar"],
        [
            ("Balance por evento", "Total contratado, pagos recibidos, gastos asociados, pago de personal y ganancia final."),
            ("Balance mensual", "Ingresos cobrados, gastos generales, pago de personal, garantias, deudas y utilidad aproximada."),
            ("Deudas pendientes", "Clientes con saldo pendiente, fecha del evento y monto por cobrar."),
            ("Gastos por categoria", "Agua, luz, impuestos, mantenimiento, publicidad, compras y otros."),
            ("Rendimiento por piso", "Cantidad de eventos, ingresos y ocupacion de cada piso."),
            ("Paquetes mas vendidos", "Ranking de paquetes, ingresos generados y frecuencia de contratacion."),
        ],
        widths=[4.2, 11.5],
    )
    add_note(
        doc,
        "Formula base",
        "Ganancia del evento = pagos recibidos del cliente - pagos al personal - gastos asociados al evento."
    )

    add_heading(doc, "12. Fases de desarrollo", 1)
    add_table(
        doc,
        ["Fase", "Objetivo", "Incluye"],
        [
            ("Fase 1", "Control principal", "Login, clientes, pisos, paquetes, calendario, eventos, contratos, pagos, boletas internas, garantias y saldo pendiente."),
            ("Fase 2", "Operacion del evento", "Checklist, personal, asignacion de personal, servicios incluidos, servicios adicionales y pago de personal."),
            ("Fase 3", "Finanzas y balance", "Gastos del local, gastos por evento, balance por evento, balance mensual, reportes y deudas."),
            ("Fase 4", "Inventario", "Articulos del local, costo, cantidad, ubicacion por piso, estado, danos, perdidas y reparaciones."),
            ("Fase 5", "Galeria y marketing", "Fotos, videos, reels, catalogo digital de paquetes y material para redes o pagina publica futura."),
        ],
        widths=[2.3, 4.2, 9.2],
    )

    add_heading(doc, "13. Prioridad recomendada", 1)
    add_table(
        doc,
        ["Prioridad", "Modulo", "Motivo"],
        [
            ("1", "Calendario por piso", "Evita cruces de fechas y permite controlar los 3 ambientes."),
            ("2", "Clientes y reservas", "Ordena la informacion de quien separa o contrata."),
            ("3", "Contratos", "Formaliza el alquiler del local y protege al negocio."),
            ("4", "Pagos y boletas internas", "Controla adelantos, saldos y movimientos de caja."),
            ("5", "Garantias", "Evita confundir dinero retenido con ingreso real."),
            ("6", "Balance basico", "Muestra rapidamente cuanto se gano o falta cobrar."),
            ("7", "Personal y gastos", "Permite conocer el costo real de operar cada evento."),
            ("8", "Inventario y reels", "Mejora control interno y apoyo comercial."),
        ],
        widths=[2.5, 4.6, 8.6],
    )

    add_heading(doc, "14. Puntos por confirmar con el cliente", 1)
    add_bullets(
        doc,
        [
            "Telefono oficial que debe aparecer en contratos y documentos.",
            "Direccion oficial del local que debe usarse en contratos y paquetes.",
            "Capacidad real de cada piso cuando el tercer piso este terminado.",
            "Si se permitiran eventos simultaneos en diferentes pisos.",
            "Monto o porcentaje oficial de separacion por tipo de paquete.",
            "Reglas de devolucion de garantia y motivos de retencion.",
            "Si las boletas internas solo seran control interno o si tambien se relacionaran con comprobantes SUNAT.",
            "Si el sistema debe generar contratos en PDF para imprimir o enviar por WhatsApp.",
            "Usuarios que tendran acceso y permisos de cada uno.",
        ],
    )

    add_heading(doc, "15. Resumen final para el cliente", 1)
    add_body(
        doc,
        "El sistema propuesto permitira que Recepciones Areli controle sus eventos de forma completa y ordenada. "
        "Cada reserva tendra cliente, piso, horario, paquete, contrato, pagos, garantia, personal, gastos y balance."
    )
    add_body(
        doc,
        "La primera version debe centrarse en fechas, pisos, contratos y dinero, porque son las areas mas importantes para evitar errores y perdida de informacion. "
        "Luego se agregaran modulos de personal, gastos, inventario y galeria comercial."
    )
    add_body(
        doc,
        "Enfoque recomendado: construir primero un sistema interno administrativo. Mas adelante se puede conectar con una pagina web publica, catalogo digital o envio por WhatsApp."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Recepciones Areli - Planificacion del Sistema Administrativo", color=MUTED, size=8.5)

    doc.save(OUT)


if __name__ == "__main__":
    main()
