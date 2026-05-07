from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "Contrato_Desarrollo_Sistema_Recepciones_Areli.docx"

BRAND = RGBColor(92, 28, 68)
DARK = RGBColor(35, 35, 35)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "EDE3EA"
NOTE_FILL = "FBF6EA"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8D0D6", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for name, size in (("Heading 1", 14), ("Heading 2", 11.5)):
        st = styles[name]
        st.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), st.font.name)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BRAND
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(4)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("CONTRATO DE PRESTACIÓN DE SERVICIOS DE DESARROLLO DE SOFTWARE")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema Administrativo para Recepciones Areli")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fecha de inicio: 05/05/2026")
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False if widths else True
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = BRAND
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_FILL)
    set_cell_border(cell, color="E3D2A7")
    set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    doc.add_paragraph()


def add_signature_block(doc):
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    labels = [
        ("EL CONTRATANTE", "EL PRESTADOR"),
        ("Firma: ______________________________", "Firma: ______________________________"),
        ("Nombre: _____________________________", "Ing. Carlos Ascarate"),
        ("DNI/RUC: ____________________________", "DNI: ________________________________"),
    ]
    for r, row in enumerate(labels):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            set_cell_border(cell, color="FFFFFF", size="0")
            set_cell_margins(cell, top=80, start=180, bottom=80, end=180)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            if r == 0:
                run.bold = True
                run.font.color.rgb = BRAND
    doc.add_paragraph()


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_heading(doc, "Primera: Partes contratantes", 1)
    add_para(
        doc,
        "Conste por el presente documento el contrato de prestación de servicios de desarrollo de software que celebran, de una parte, "
        "el/la representante de Recepciones Areli, a quien en adelante se le denominará EL CONTRATANTE; y de la otra parte, "
        "el Ing. Carlos Ascarate, a quien en adelante se le denominará EL PRESTADOR.",
    )
    add_para(
        doc,
        "Los datos completos de EL CONTRATANTE serán consignados al momento de la firma: nombre o razón social, DNI/RUC, domicilio, teléfono y correo de contacto.",
    )

    add_heading(doc, "Segunda: Objeto del contrato", 1)
    add_para(
        doc,
        "EL CONTRATANTE encarga a EL PRESTADOR el análisis, diseño, desarrollo e implementación de un Sistema Administrativo para Recepciones Areli, "
        "orientado a controlar reservas, contratos, pagos, paquetes, personal, gastos, inventario, galería de eventos y balance financiero del negocio.",
    )

    add_heading(doc, "Tercera: Alcance funcional del sistema", 1)
    add_para(doc, "El sistema comprenderá, como alcance principal, los siguientes módulos y funcionalidades:")
    add_bullets(
        doc,
        [
            "Gestión de usuarios y acceso al sistema.",
            "Registro y administración de clientes.",
            "Gestión de pisos o ambientes: 1er piso, 2do piso y 3er piso.",
            "Calendario de reservas por fecha, horario, piso y estado del evento.",
            "Gestión de paquetes de eventos: fiestas, matrimonios, 15 años, promociones escolares y paquetes personalizados.",
            "Registro de contratos con datos del cliente, evento, piso, horario, monto, modalidad de pago y condiciones.",
            "Registro de pagos de clientes, adelantos, pagos parciales, saldo pendiente y medio de pago.",
            "Generación o registro de boletas internas de pago.",
            "Control de garantías recibidas, devueltas o retenidas por desperfectos.",
            "Organización del evento por personal: encargada, mozos, barman, DJ, animador, seguridad, limpieza y otros roles necesarios.",
            "Control de pago de personal por evento.",
            "Registro de gastos del local: agua, luz, SUNAT, impuestos, mantenimiento, compras, publicidad y otros.",
            "Inventario del local con cantidad, estado, ubicación, precio unitario y costo total.",
            "Galería o registro de fotos, videos y reels de eventos.",
            "Balance financiero por evento y reportes básicos de ingresos, egresos, saldos pendientes y utilidad estimada.",
        ],
    )

    add_heading(doc, "Cuarta: Plazo de ejecución", 1)
    add_para(
        doc,
        "El plazo estimado para la ejecución del servicio será de dos meses y medio, iniciando el día 05/05/2026 y culminando aproximadamente el día 20/07/2026, "
        "siempre que EL CONTRATANTE entregue oportunamente la información, documentos, validaciones y observaciones necesarias para avanzar con el proyecto.",
    )

    add_heading(doc, "Quinta: Monto y forma de pago", 1)
    add_para(
        doc,
        "El monto total pactado por el servicio asciende a S/ 6,000.00 (seis mil con 00/100 soles).",
    )
    add_table(
        doc,
        ["Concepto", "Monto"],
        [
            ("Monto total del proyecto", "S/ 6,000.00"),
            ("Adelanto a la firma del contrato", "S/ 1,500.00"),
            ("Saldo pendiente", "S/ 4,500.00"),
            ("Forma de pago del saldo", "10 pagos semanales de S/ 450.00"),
        ],
        widths=[8.0, 7.0],
    )
    add_para(
        doc,
        "El adelanto de S/ 1,500.00 será pagado al inicio del proyecto. El saldo restante de S/ 4,500.00 será cancelado en diez pagos semanales de S/ 450.00 cada uno.",
    )

    add_heading(doc, "Sexta: Cronograma de pagos", 1)
    add_table(
        doc,
        ["Pago", "Fecha referencial", "Monto", "Estado"],
        [
            ("Adelanto", "05/05/2026", "S/ 1,500.00", "Pendiente de confirmar"),
            ("Semana 1", "12/05/2026", "S/ 450.00", "Pendiente"),
            ("Semana 2", "19/05/2026", "S/ 450.00", "Pendiente"),
            ("Semana 3", "26/05/2026", "S/ 450.00", "Pendiente"),
            ("Semana 4", "02/06/2026", "S/ 450.00", "Pendiente"),
            ("Semana 5", "09/06/2026", "S/ 450.00", "Pendiente"),
            ("Semana 6", "16/06/2026", "S/ 450.00", "Pendiente"),
            ("Semana 7", "23/06/2026", "S/ 450.00", "Pendiente"),
            ("Semana 8", "30/06/2026", "S/ 450.00", "Pendiente"),
            ("Semana 9", "07/07/2026", "S/ 450.00", "Pendiente"),
            ("Semana 10", "14/07/2026", "S/ 450.00", "Pendiente"),
        ],
        widths=[3.0, 4.2, 3.4, 4.4],
    )

    add_heading(doc, "Séptima: Obligaciones de EL PRESTADOR", 1)
    add_bullets(
        doc,
        [
            "Realizar el análisis y desarrollo del sistema conforme al alcance acordado.",
            "Presentar avances del proyecto para revisión de EL CONTRATANTE.",
            "Corregir errores detectados durante el desarrollo que correspondan al alcance contratado.",
            "Mantener comunicación sobre avances, bloqueos y requerimientos de información.",
            "Entregar el sistema funcional conforme a las etapas acordadas.",
        ],
    )

    add_heading(doc, "Octava: Obligaciones de EL CONTRATANTE", 1)
    add_bullets(
        doc,
        [
            "Entregar información, documentos, formatos, datos de paquetes, reglas de negocio y cualquier material necesario para el desarrollo.",
            "Revisar los avances y responder observaciones en tiempos razonables.",
            "Realizar los pagos conforme al cronograma acordado.",
            "Designar una persona responsable para validar funcionalidades y contenidos.",
            "Informar oportunamente cualquier cambio importante en el alcance del sistema.",
        ],
    )

    add_heading(doc, "Novena: Cambios de alcance", 1)
    add_para(
        doc,
        "Cualquier funcionalidad adicional que no se encuentre descrita en el alcance de este contrato será evaluada por ambas partes. "
        "Si el cambio implica mayor tiempo de desarrollo, podrá generar una ampliación de plazo y/o un costo adicional previamente acordado.",
    )

    add_heading(doc, "Décima: Entrega y conformidad", 1)
    add_para(
        doc,
        "La entrega del sistema se realizará de forma progresiva, por módulos o etapas. EL CONTRATANTE deberá revisar y validar los avances presentados. "
        "La conformidad final se dará cuando el sistema cumpla con el alcance principal acordado en este contrato.",
    )

    add_heading(doc, "Décima primera: Suspensión por incumplimiento de pago", 1)
    add_para(
        doc,
        "En caso de retraso injustificado en los pagos pactados, EL PRESTADOR podrá suspender temporalmente el avance del proyecto hasta que se regularice el pago pendiente. "
        "Los retrasos de pago podrán afectar el cronograma de entrega.",
    )

    add_heading(doc, "Décima segunda: Confidencialidad", 1)
    add_para(
        doc,
        "Ambas partes se comprometen a mantener reserva sobre información interna, datos de clientes, precios, documentos, accesos, reportes o cualquier información sensible compartida durante el desarrollo del proyecto.",
    )

    add_heading(doc, "Décima tercera: Propiedad y uso del sistema", 1)
    add_para(
        doc,
        "Una vez cancelado el monto total pactado, EL CONTRATANTE tendrá derecho de uso del sistema desarrollado para la administración interna de Recepciones Areli. "
        "Los componentes técnicos, librerías, plantillas, estructuras reutilizables o conocimientos generales utilizados por EL PRESTADOR podrán ser reutilizados en otros proyectos, siempre que no se divulgue información confidencial de EL CONTRATANTE.",
    )

    add_heading(doc, "Décima cuarta: Aceptación", 1)
    add_para(
        doc,
        "Leído el presente contrato y estando ambas partes conformes con su contenido, lo firman en señal de aceptación, en la ciudad de Lima, con fecha 05/05/2026.",
    )

    add_signature_block(doc)

    add_note(
        doc,
        "Nota: Este documento es un modelo contractual operativo para formalizar el acuerdo entre las partes. Puede ser revisado o complementado con asesoría legal si las partes lo consideran necesario.",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Contrato de Desarrollo de Software - Recepciones Areli")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    doc.save(OUT)


if __name__ == "__main__":
    main()
