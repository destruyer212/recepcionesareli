from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(r"C:\SpringProjectsnew\recepcionesareli")
SOURCE_LOGO = Path(r"C:\Users\CARLOS\Downloads\img20.png")
OUTPUT = ROOT / "documentacion" / "word" / "Contrato_Areli_Piso_3_4_Pack_VIP_Pro_Final.docx"
WATERMARK = ROOT / "documentacion" / "metadata" / "areli_watermark_soft.png"

BLUE = RGBColor(31, 78, 121)
GOLD = RGBColor(196, 145, 0)
LIGHT_GOLD = "FFF4D6"
BORDER_GOLD = "B99952"
TEXT = RGBColor(30, 30, 30)


def make_watermark() -> None:
    WATERMARK.parent.mkdir(parents=True, exist_ok=True)
    image = Image.open(SOURCE_LOGO).convert("RGBA")
    image.thumbnail((760, 620))
    alpha = image.getchannel("A").point(lambda value: int(value * 0.16))
    image.putalpha(alpha)
    canvas = Image.new("RGBA", (900, 720), (255, 255, 255, 0))
    canvas.alpha_composite(image, ((canvas.width - image.width) // 2, (canvas.height - image.height) // 2))
    canvas.save(WATERMARK)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER_GOLD, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")


def style_run(run, *, bold=False, size=10, color=TEXT) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", *, bold=False, size=10, align=None, color=TEXT, spacing=3):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(spacing)
    paragraph.paragraph_format.line_spacing = 1.05
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    style_run(run, bold=bold, size=size, color=color)
    return paragraph


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    style_run(run, bold=True, size=11, color=RGBColor(0, 0, 0))


def add_clause(doc: Document, title: str, paragraphs: list[str]) -> None:
    add_heading(doc, title)
    for text in paragraphs:
        paragraph = add_paragraph(doc, text, size=9.3, spacing=3)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style=None)
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.first_line_indent = Cm(-0.25)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.02
    run = paragraph.add_run("- " + text)
    style_run(run, size=9.2)


def add_info_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (7.8, 8.2)) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table)
    for row_idx, (label, value) in enumerate(rows):
        row = table.rows[row_idx]
        row.cells[0].width = Cm(widths[0])
        row.cells[1].width = Cm(widths[1])
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], LIGHT_GOLD)
        label_run = row.cells[0].paragraphs[0].add_run(label)
        style_run(label_run, bold=True, size=9.5, color=RGBColor(0, 0, 0))
        value_run = row.cells[1].paragraphs[0].add_run(value)
        style_run(value_run, size=9.5, color=RGBColor(0, 0, 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_signature_table(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=2)
    set_table_width(table)
    data = [
        ("______________________________", "______________________________"),
        ("EL ARRENDADOR", "EL ARRENDATARIO"),
        ("DNI: {{arrendador_dni}}", "DNI: {{cliente_dni}}"),
        ("Firma: {{firma_arrendador}}", "Firma: {{firma_cliente}}"),
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, color="303030", size="6")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            style_run(run, bold=row_idx == 1, size=9.5 if row_idx != 1 else 10.5, color=RGBColor(0, 0, 0))


def add_logo(doc: Document, image_path: Path, *, width_inches: float, spacing_after: int = 2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def add_header_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Documento generado por sistema - Codigo: {{codigo_contrato}}")
    style_run(run, size=8, color=RGBColor(95, 95, 95))


def build_document() -> None:
    make_watermark()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.3)
    section.footer_distance = Cm(0.35)
    add_header_footer(section)

    add_logo(doc, SOURCE_LOGO, width_inches=1.25, spacing_after=1)
    add_paragraph(doc, "ARELI", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, spacing=0)
    add_paragraph(doc, "SALON DE RECEPCIONES", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, spacing=1)
    add_paragraph(doc, "CONTRATO DE ALQUILER DEL LOCAL", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0, 0, 0), spacing=2)
    add_paragraph(doc, "SALON DE RECEPCIONES ARELI", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=GOLD, spacing=1)
    add_paragraph(doc, "3ER Y 4TO PISO - PACK VIP / PAQUETE CONTRATADO", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=GOLD, spacing=8)

    add_info_table(doc, [
        ("Codigo de contrato", "{{codigo_contrato}}"),
        ("Fecha de emision", "{{fecha_emision}}"),
    ])

    intro = (
        "Conste por el presente documento el contrato de alquiler temporal de local para evento privado que celebran, "
        "de una parte, {{arrendador_nombre}}, identificado(a) con DNI N. {{arrendador_dni}}, con domicilio en "
        "{{arrendador_direccion}}, a quien en adelante se le denominara EL ARRENDADOR; y de la otra parte, "
        "{{cliente_nombre}}, identificado(a) con DNI N. {{cliente_dni}}, con domicilio en {{cliente_direccion}} "
        "y celular {{cliente_celular}}, a quien en adelante se le denominara EL ARRENDATARIO; bajo los terminos "
        "y condiciones siguientes:"
    )
    add_paragraph(doc, intro, size=9.4, spacing=5).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_heading(doc, "DATOS GENERALES DEL EVENTO")
    add_info_table(doc, [
        ("Tipo de evento", "{{tipo_evento}}"),
        ("Fecha del evento", "{{fecha_evento}}"),
        ("Horario contratado", "Desde {{hora_inicio}} hasta {{hora_fin}} del dia {{fecha_fin}}"),
        ("Ambientes contratados", "{{ambientes_contratados}}"),
        ("Paquete contratado", "{{paquete_nombre}}"),
        ("Capacidad maxima permitida", "{{capacidad_maxima}} personas"),
        ("Monto total del contrato", "S/ {{monto_total}}"),
        ("Adelanto / separacion", "S/ {{adelanto}}"),
        ("Saldo pendiente", "S/ {{saldo}}"),
        ("Garantia", "S/ {{garantia_danos}}"),
        ("Pago APDAYC", "S/ {{apdayc_monto}} - Asume: {{apdayc_responsable}} - Estado: {{apdayc_estado}}"),
        ("Medio de pago", "{{medio_pago}}"),
    ], widths=(6.5, 9.5))

    add_clause(doc, "CLAUSULA PRIMERA: OBJETO DEL CONTRATO", [
        "EL ARRENDADOR entrega en alquiler temporal a EL ARRENDATARIO el uso del ambiente indicado en los datos generales, correspondiente al local denominado Salon de Recepciones Areli, ubicado en Calle Los Nisperos Mz. B1 Lote 19, distrito de Carabayllo, para la realizacion exclusiva del evento declarado.",
        "La entrega del ambiente se efectuara en condiciones operativas para el desarrollo del evento, de acuerdo con la contratacion real, disponibilidad del local y servicios expresamente pactados."
    ])

    add_clause(doc, "CLAUSULA SEGUNDA: PAQUETE Y SERVICIOS CONTRATADOS", [
        "EL ARRENDATARIO contrata el paquete {{paquete_nombre}}. Los servicios, alcances, inclusiones y condiciones especificas del paquete seleccionado se detallan en el ANEXO FINAL del presente contrato, el cual forma parte integrante del documento.",
        "Cualquier servicio no indicado expresamente en el paquete, anexo, comprobante o acuerdo escrito se considerara adicional y podra generar un costo separado."
    ])

    add_clause(doc, "CLAUSULA TERCERA: MONTO, ADELANTO, SALDO Y REGISTRO DE PAGOS", [
        "El monto total pactado es de S/ {{monto_total}}. EL ARRENDATARIO declara haber entregado como adelanto o separacion la suma de S/ {{adelanto}}, quedando un saldo pendiente de S/ {{saldo}}, el cual debera ser cancelado como maximo hasta {{fecha_limite_pago}}, salvo acuerdo escrito distinto entre las partes.",
        "La reserva del local queda confirmada con el pago de la separacion correspondiente. Todo pago debe quedar respaldado mediante recibo, constancia, voucher, comprobante o registro generado por el sistema administrativo de Recepciones Areli."
    ])

    add_clause(doc, "CLAUSULA CUARTA: GARANTIA Y CONDICIONES DE NO DEVOLUCION POR DESISTIMIENTO", [
        "La garantia de S/ {{garantia_danos}} respalda la reserva, disponibilidad del ambiente, coordinaciones previas, gastos administrativos y eventuales danos, perdidas, deterioros, saldos u observaciones generadas durante el evento.",
        "En caso EL ARRENDATARIO desista unilateralmente del evento, cancele por decision propia o no continue con la contratacion, los importes entregados por separacion, adelanto y/o garantia no seran materia de devolucion, por encontrarse vinculados a la reserva de fecha, bloqueo de disponibilidad del local y gestiones realizadas por EL ARRENDADOR.",
        "Esta condicion se establece sin perjuicio de que EL ARRENDADOR pueda evaluar, de manera excepcional y por escrito, una reprogramacion conforme a las reglas del presente contrato."
    ])

    add_clause(doc, "CLAUSULA QUINTA: APDAYC, IGV, DERECHOS Y PERMISOS NO INCLUIDOS", [
        "El pago de APDAYC, IGV, derechos de autor, permisos municipales, autorizaciones especiales u otros conceptos no incluidos expresamente en el paquete sera asumido por EL ARRENDATARIO, promocion, institucion contratante o responsable del evento, salvo acuerdo escrito distinto.",
        "Para este contrato, el concepto APDAYC se registra como: monto S/ {{apdayc_monto}}, responsable {{apdayc_responsable}}, estado {{apdayc_estado}}. Observacion: {{apdayc_observacion}}."
    ])

    add_clause(doc, "CLAUSULA SEXTA: HORARIO, ENTREGA Y HORA EXTRA", [
        "EL ARRENDATARIO se obliga a respetar el horario contratado. Si el evento inicia en la noche y finaliza de madrugada, se entendera que la hora de termino corresponde al dia calendario siguiente, siempre que asi figure en los datos generales.",
        "Al llegar la hora de finalizacion, los asistentes deberan desocupar el local para permitir la revision, cierre y entrega del ambiente. En caso de exceder el horario pactado, EL ARRENDATARIO debera pagar S/ {{costo_hora_extra}} por cada hora adicional o fraccion, siempre que EL ARRENDADOR autorice expresamente la extension."
    ])

    add_clause(doc, "CLAUSULA SEPTIMA: RESPONSABILIDAD POR DANOS, DETERIOROS Y BIENES PERSONALES", [
        "EL ARRENDATARIO sera responsable por cualquier desperfecto, deterioro, perdida, rotura o dano ocasionado al local, mobiliario, paredes, pisos, puertas, instalaciones electricas, servicios higienicos, equipos, accesorios o bienes entregados para el uso del evento.",
        "EL ARRENDADOR no se responsabiliza por perdidas, robos, accidentes fortuitos, objetos extraviados, conflictos entre asistentes o hechos ocasionados por terceros dentro o fuera del local. La seguridad de menores de edad sera responsabilidad de sus padres, tutores o adultos responsables."
    ])

    add_clause(doc, "CLAUSULA OCTAVA: PROHIBICIONES", [
        "Queda expresamente prohibido a EL ARRENDATARIO realizar actos que afecten la seguridad, conservacion, orden, reputacion o correcto funcionamiento del local."
    ])
    for item in [
        "Subarrendar, ceder o transferir el uso del local a terceros sin autorizacion escrita.",
        "Exceder la capacidad maxima permitida.",
        "Realizar actividades distintas al evento declarado.",
        "Danar, perforar, pintar, modificar o alterar paredes, pisos, puertas, instalaciones o accesorios del local.",
        "Usar pirotecnicos, materiales inflamables, elementos peligrosos o equipos no autorizados.",
        "Permanecer en el local fuera del horario contratado sin autorizacion expresa.",
    ]:
        add_bullet(doc, item)

    add_clause(doc, "CLAUSULA NOVENA: CAMBIO DE FECHA Y REPROGRAMACION", [
        "Toda solicitud de cambio de fecha debe comunicarse por escrito con una anticipacion minima de quince (15) dias calendario antes de la fecha del evento. Por ejemplo, si el evento se encuentra programado para el 30 de junio, la solicitud debera presentarse como maximo hasta el 15 de junio.",
        "Si EL ARRENDADOR acepta la solicitud, EL ARRENDATARIO tendra un plazo maximo de dos (2) meses para fijar una nueva fecha, siempre sujeto a disponibilidad del local, fechas libres, terminos vigentes, condiciones operativas y aprobacion expresa de EL ARRENDADOR.",
        "Si vencido dicho plazo no se define una nueva fecha aceptada por EL ARRENDADOR, o si EL ARRENDATARIO no cumple las condiciones de reprogramacion, se perdera la garantia y los importes asociados a la separacion, sin derecho a reclamo posterior."
    ])

    add_clause(doc, "CLAUSULA DECIMA: ACTA DE ENTREGA Y CONFORMIDAD", [
        "Al finalizar el evento, ambas partes podran verificar el estado de entrega del local. De existir observaciones, estas podran registrarse en acta, fotografias, videos o anotacion del sistema, indicando danos, perdidas, saldos pendientes u otras incidencias."
    ])

    add_clause(doc, "CLAUSULA DECIMO PRIMERA: ACEPTACION", [
        "Ambas partes declaran haber leido el presente contrato, comprender su contenido y estar conformes con todas sus clausulas, firmandolo en senal de aceptacion en la ciudad de Carabayllo, a los {{dia_firma}} dias del mes de {{mes_firma}} de {{anio_firma}}."
    ])
    add_signature_table(doc)
    add_paragraph(doc, "Observacion del sistema: {{observacion_contrato}}", size=8.5, color=RGBColor(95, 95, 95), spacing=3)

    doc.add_page_break()
    add_logo(doc, WATERMARK, width_inches=1.05, spacing_after=1)
    add_paragraph(doc, "ANEXO FINAL", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, color=GOLD, spacing=2)
    add_paragraph(doc, "DETALLE DEL PAQUETE CONTRATADO", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, spacing=8)
    add_info_table(doc, [
        ("Paquete seleccionado", "{{paquete_nombre}}"),
        ("Ambiente", "{{ambientes_contratados}}"),
        ("Capacidad referencial", "{{capacidad_maxima}} personas"),
        ("Precio del paquete", "S/ {{monto_total}}"),
        ("Garantia", "S/ {{garantia_danos}}"),
        ("APDAYC / derechos", "{{apdayc_responsable}} - {{apdayc_estado}} - S/ {{apdayc_monto}}"),
    ], widths=(6.2, 9.8))

    add_heading(doc, "SERVICIOS INCLUIDOS")
    add_paragraph(doc, "{{detalle_paquete_contratado}}", size=9.4, spacing=5).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_heading(doc, "CONDICIONES DEL PAQUETE")
    add_paragraph(doc, "{{condiciones_paquete}}", size=9.4, spacing=5).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_heading(doc, "NOTA OPERATIVA")
    add_paragraph(
        doc,
        "Este anexo se completa automaticamente segun el paquete elegido en el sistema: Basico, Areli, Premium, Promociones escolares o un paquete personalizado. Si el paquete seleccionado no incluye APDAYC, IGV, permisos u otros derechos, dichos conceptos seran asumidos conforme a lo indicado en la clausula correspondiente.",
        size=9.4,
        spacing=5,
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
